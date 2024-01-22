import os
import sys
from re import S
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import SentenceTransformerEmbeddings, HuggingFaceBgeEmbeddings
import pinecone
from langchain_community.vectorstores import Pinecone
import json
from PyPDF2 import PdfReader
import numpy
import argparse
from docx import Document as DocumentReader
from docx.opc.exceptions import PackageNotFoundError


class Document:
    def __init__(self, page_content, metadata=None):
        self.page_content = page_content
        self.metadata = metadata if metadata is not None else {}


def split_docs(documents,chunk_size=500,chunk_overlap=20):
    """
    Splits the documents into chunks.
    :param documents: List with a single element containing text of document
    :type documents: list[Document]
    :param chunk_size: Maximum number of characters per chunk. Default 500.
    :type chunk_size: int
    :param chunk_overlap: Number of characters each chunk overlaps the previous chunk.
    :type chunk_overlap: int
    :return: A list of chunks.
    :rtype: list[Document]
    """
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    docs = text_splitter.split_documents(documents)
    return docs
  

def read_docx(file_path):
    """
    Extrapolates text from a .docx file type.
    :param file_path: The relative file path of target document.
    :type file_path: str
    :return: The text of the document.
    :rtype: str
    """
    document = DocumentReader(file_path)
    text = ""
    for paragraph in document.paragraphs:
        text += paragraph.text + "\n"
    return text


def main():
    # Load shell configuration
    parser = argparse.ArgumentParser(description="Embed and index document into defined namespace.")
    parser.add_argument('-n', '--namespace', help='Specify the Pinecone namespace.', required=True)
    parser.add_argument('-d', '--debug', help='Provides additional information in output.', action='store_true')
    parser.add_argument('-f', '--filename', help='The name of the file.', required=True)
    args = parser.parse_args()

    # Load configuration
    print("\n(1/6) Loading configuration...")
    with open('../appSettings.Development.json') as parser:
      appSettings = json.load(parser)
    if args.debug:
        print(appSettings)

    # Load documents
    print("\n(2/6) Loading document(s)...")
    documents = []
    documentFileName = args.filename
    documentFilePath = f'../UserFiles/{documentFileName}'
    try:
        if ".pdf" in documentFilePath:
            reader = PdfReader(documentFilePath)
            page = reader.pages[0] 
            text = page.extract_text() 
            documents.append(Document(page_content=text))
        if args.debug:
            print(documents)
        elif ".docx" in documentFilePath:
            text = read_docx(documentFilePath)
            documents.append(Document(page_content=text))
            if args.debug:
                print(documents)
        else:
            print("File type not supported.")
            return
    except PackageNotFoundError as e:
        print("File does not exist.")
        return
    except FileNotFoundError as e:
        print("File does not exist.")
        return

    # Split document(s)
    print("\n(3/6) Splitting document(s)...")
    docs = split_docs(documents)
    if args.debug:
        print(docs)

    # Initialize embedding model
    print("\n(4/6) Initializing embedding model...")
    model_name = "all-MiniLM-L6-v2"
    embeddings = HuggingFaceBgeEmbeddings(model_name=model_name)
    if args.debug:
        print("Model Name: " + model_name)
        print(embeddings)

    # Initialize Pinecone index
    print("\n(5/6) Initializing Pinecone index...")
    pinecone.init(
        api_key = appSettings.get("Pinecone", {}).get("APIKey", None),
        environment= appSettings.get("Pinecone", {}).get("Environment", None)
    )
    index_name = appSettings.get("Pinecone", {}).get("Index", None)
    namespace = args.namespace
    try:
        index = Pinecone.from_documents(docs, embeddings, index_name=index_name, namespace=namespace)
    except Exception as e:
        print(e)
    if args.debug:
        print("Index Name: " + index_name)
        print("Namespace: " + namespace)
        print("Number of documents: " + str(len(docs)))

    # Complete
    print("\n(6/6) Finished.")


if __name__ == "__main__":
    main()
